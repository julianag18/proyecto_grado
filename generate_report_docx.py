import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_report():
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("INFORME DE AVANCE PARCIAL:\nDIGITALIZACIÓN Y OPTIMIZACIÓN DEL MÓDULO METROLÓGICO (PAME)")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Proyecto de Grado para la Titulación Profesional\n").bold = True
    meta.add_run("Autor: ").bold = True
    meta.add_run("Juliana Gómez\n")
    meta.add_run("Fecha de Generación: ").bold = True
    meta.add_run("29 de Julio de 2026\n")
    meta.add_run("Destinatario: ").bold = True
    meta.add_run("Asesor Interno del Proyecto de Grado / Profesor de Universidad\n")
    
    # Divider
    p = doc.add_paragraph()
    p.add_run("―" * 45).font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Heading 1 helper
    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        r = h.add_run(text)
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E) # Teal color matching PAME theme
        return h

    # Heading 2 helper
    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        r = h.add_run(text)
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x11, 0x5E, 0x59)
        return h

    # Paragraph helper
    def add_para(text, before=0, after=6):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(before)
        para.paragraph_format.space_after = Pt(after)
        para.paragraph_format.line_spacing = 1.15
        r = para.add_run(text)
        return r

    # Section 1
    add_heading_1("1. Resumen Ejecutivo")
    add_para(
        "El presente informe de avance parcial expone de manera detallada el diseño, la reestructuración "
        "y el estado técnico del módulo digital de Plan de Aseguramiento Metrológico (PAME) concebido para "
        "Laboratorios Laproff S.A.S. En el transcurso de las últimas semanas, se abordó la necesidad de "
        "migrar un cronograma tradicional en hojas de cálculo hacia una plataforma web centralizada, interactiva "
        "y de alto rendimiento. "
    )
    add_para(
        "A través de una reingeniería de bases de datos, se resolvió un cuello de botella de rendimiento que congelaba "
        "el software durante varios minutos ante cargas de datos masivas. Asimismo, se integró el motor de notificaciones "
        "automáticas por correo electrónico a través de la API SMTP de Brevo. Este informe describe detalladamente el "
        "CÓMO y el PORQUÉ de las decisiones arquitectónicas tomadas, sirviendo de base para la retroalimentación del asesor "
        "antes de proceder a la redacción formal de la tesis y las validaciones en sitio."
    )

    # Section 2
    add_heading_1("2. Arquitectura del Sistema: ¿Cómo y Con Qué se Construyó?")
    add_para(
        "Para garantizar un desarrollo ágil, seguro y mantenible, se seleccionó un stack tecnológico robusto, "
        "orientado a la representación de datos analíticos en tiempo real:"
    )
    
    tech_points = [
        ("Interfaz de Usuario (Frontend): ", 
         "Desarrollada en Streamlit (Python) acoplado a un sistema de diseño web premium mediante CSS nativo. Esto proporciona una navegación fluida basada en pestañas (Dashboard, Inventario, Cumplimiento, Cronograma y Migración ETL) y un diseño responsivo adaptado a dispositivos móviles y estaciones de cómputo en planta."),
        ("Base de Datos (Persistencia): ", 
         "Implementada sobre Firebase / Google Cloud Firestore en modo NoSQL y adaptada con repositorios locales parametrizables (Modo Demo) para garantizar portabilidad e independencia de infraestructura durante el desarrollo preliminar."),
        ("Motor de Alertas y Notificaciones (Integración Externa): ", 
         "Conectado a través del protocolo SMTP cifrado con Brevo (Sendinblue). El sistema despacha correos automáticos estructurados en HTML responsivo que contienen tablas analíticas e indicadores clave de rendimiento (KPIs) en tiempo real.")
    ]
    
    for title, desc in tech_points:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run_title = p.add_run(title)
        run_title.bold = True
        run_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        p.add_run(desc)

    # Section 3
    add_heading_1("3. Decisiones Técnicas Clave y Justificaciones (El \"Por Qué\")")
    
    add_heading_2("A. Eliminación del congelamiento del sistema (Optimización del Backend)")
    add_para(
        "Desafío Detectado: Al cargar la base de datos completa de Laboratorios Laproff, que incluye cientos de registros de "
        "equipos y múltiples servicios históricos de calibración/validación, la aplicación sufría un congelamiento de hasta 3 minutos "
        "al cambiar de pestaña. El análisis del log arrojó un error de diseño de tipo N+1: el sistema realizaba consultas consecutivas "
        "a la base de datos por cada celda y equipo renderizado en pantalla."
    )
    add_para(
        "Solución Implementada: Se reescribió el repositorio de datos para descargar la totalidad de los servicios en una sola consulta "
        "agrupada, resolviendo la relación en memoria con complejidad de tiempo O(N) lineal. Adicionalmente, se decoraron las funciones "
        "de lectura con almacenamiento en caché local (@st.cache_data) y se configuró un mecanismo automático de invalidación que "
        "borra la caché solo cuando se ejecuta una nueva migración de datos, asegurando que la información permanezca al día sin "
        "saturar el servidor. El paso entre pestañas pasó a ser instantáneo (milisegundos)."
    )

    add_heading_2("B. Reglas de Envío de Alertas por Lotes para Evitar Fatiga por Notificaciones")
    add_para(
        "Desafío Detectado: Enviar correos electrónicos diarios por cada equipo próximo a vencer genera saturación ("
        "fatiga por alertas) en la bandeja de entrada del coordinador de metrología, lo que usualmente conduce a que las notificaciones "
        "sean ignoradas."
    )
    add_para(
        "Solución Implementada: Se programó una regla de negocio inteligente. Las alertas estándar (equipos en estado 'Programar' con 15 a 45 "
        "días restantes) se retienen y solo se despachan automáticamente una vez acumuladas en un lote consolidado de cinco (5) o más "
        "equipos. No obstante, si el sistema detecta algún equipo en estado 'Crítico' (menos de 15 días para vencer), la regla del lote se "
        "evade automáticamente y se despacha una alerta roja de forma inmediata."
    )

    # Section 4
    add_heading_1("4. Justificación Metrológica del Plazo de Alerta de 1 Mes (30 días)")
    add_para(
        "Uno de los puntos clave a defender ante el jurado y el asesor universitario es la selección del plazo preventivo de "
        "30 días para las alertas automáticas. En metrología industrial y manufactura farmacéutica, el vencimiento de un instrumento "
        "implica su retiro inmediato del proceso productivo, lo que puede detener líneas enteras de envasado, dosificación o control de calidad. "
        "Por ende, 30 días es el margen óptimo debido al siguiente ciclo logístico real:"
    )

    bullet_points = [
        ("Trámite Administrativo y Cotización (Días 1 a 12): ", 
         "El coordinador debe documentar las especificaciones y tolerancias del instrumento, solicitar cotizaciones a proveedores externos que cuenten con acreditación ONAC (u homólogos vigentes) y tramitar la aprobación del gasto con el departamento de compras."),
        ("Programación Operativa (Días 12 a 17): ", 
         "Se negocia con el área de producción del laboratorio para hallar ventanas de tiempo en las que el equipo pueda calibrarse en sitio o enviarse al laboratorio del proveedor, minimizando el impacto en la cadena de manufactura de medicamentos."),
        ("Ejecución Técnica del Servicio (Días 17 a 24): ", 
         "Corresponde al traslado físico del patrón o del instrumento, ejecución del ensayo metrológico, cálculo de incertidumbres y el tiempo de emisión del informe técnico por parte del laboratorio externo."),
        ("Entrega y Dictamen de Conformidad (Días 24 a 30): ", 
         "El coordinador recibe el equipo y el certificado de calibración. Se realiza un análisis de tolerancia del proceso para verificar si la desviación del instrumento cumple con los requisitos del método analítico. Si cumple, se etiqueta como 'Conforme' y se reincorpora oficialmente a planta antes de la fecha límite.")
    ]
    
    for title, desc in bullet_points:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run_title = p.add_run(title)
        run_title.bold = True
        run_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        p.add_run(desc)

    # Section 5
    add_heading_1("5. Resumen de Pruebas y Validación Técnica")
    add_para(
        "El módulo cuenta con una suite de pruebas de caja blanca utilizando pytest en la carpeta tests/. "
        "Estas pruebas automatizadas simulan cargas masivas a la base de datos, el cálculo matemático de días restantes, "
        "las reglas lógicas de transición de estados del cronograma y el formateo dinámico de correos en HTML. "
        "Todas las pruebas integradas pasan con éxito (13 pruebas aprobadas de forma limpia), lo que garantiza la "
        "estabilidad estructural de la aplicación ante cambios futuros."
    )

    # Section 6
    add_heading_1("6. Roadmap y Puntos Clave para la Reunión con el Asesor")
    add_para(
        "A continuación se enlistan las áreas de mejora visual y técnica en las que se continuará trabajando "
        "y que servirán de base para la retroalimentación inmediata del profesor:"
    )

    roadmap_points = [
        ("Mejoras de Frontend y Visualización de Gráficos: ", "Ajustar las leyendas y títulos de los ejes X e Y de las gráficas de Plotly. En pantallas angostas, algunos nombres extensos de áreas del laboratorio tienden a recortarse."),
        ("Retroalimentación sobre Pruebas de Campo: ", "Definir el protocolo de validación y el tiempo óptimo del piloto con datos reales del laboratorio. ¿Es recomendable mantener el paralelo con el sistema anterior por 1 o 2 semanas?"),
        ("Definición de Capítulos de Tesis: ", "Presentar la estructura inicial del documento escrito de grado para recibir sus sugerencias en cuanto a los apartados teóricos de aseguramiento metrológico e ingeniería de software.")
    ]
    
    for title, desc in roadmap_points:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run_title = p.add_run(title)
        run_title.bold = True
        run_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        p.add_run(desc)

    # Save to Desktop
    import os
    saved_successfully = False
    
    desktop_paths = [
        "C:/Users/julianag18/Desktop/informe_avance_proyecto.docx",
        "C:/Users/julianag18/OneDrive/Desktop/informe_avance_proyecto.docx"
    ]
    
    for path in desktop_paths:
        try:
            # Create directory if it doesn't exist (though desktop should exist)
            os.makedirs(os.path.dirname(path), exist_ok=True)
            doc.save(path)
            print(f"Report saved to Desktop at: {path}")
            saved_successfully = True
        except Exception as e:
            print(f"Could not save to {path}: {e}")
            
    if not saved_successfully:
        # Fallback to local workspace if desktop write fails entirely
        doc.save("informe_avance_proyecto_desktop_backup.docx")
        print("Desktop write failed. Saved locally as: informe_avance_proyecto_desktop_backup.docx")

    # Clean up old versions in the project directories
    cleanup_files = [
        "C:/Users/julianag18/OneDrive/Desktop/proyecto_grado-main/informe_avance_proyecto.docx",
        "C:/Users/julianag18/OneDrive/Desktop/proyecto_grado-main/informe_avance_proyecto_v2.docx",
        "C:/Users/julianag18/Desktop/Proyecto de grado/proyecto_grado-main/informe_avance_proyecto.docx",
        "C:/Users/julianag18/Desktop/Proyecto de grado/proyecto_grado-main/informe_avance_proyecto_v2.docx",
        "informe_avance_proyecto.docx",
        "informe_avance_proyecto_v2.docx"
    ]
    for file_path in cleanup_files:
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                print(f"Deleted old version: {file_path}")
        except Exception as e:
            print(f"Could not delete old version {file_path}: {e}")

if __name__ == "__main__":
    create_report()
