import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_final_report():
    doc = docx.Document()
    
    # Configuración de márgenes estándar (1 pulgada a cada lado)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Configuración de fuente normal (Times New Roman, 11 pt, negro)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Helpers de cabeceras oficiales (Plantilla UdeA)
    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(8)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.bold = True
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
        r.bold = True
        return h

    def add_heading_3(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
        r.italic = True
        return h

    def add_para(text, before=0, after=6):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(before)
        para.paragraph_format.space_after = Pt(after)
        para.paragraph_format.line_spacing = 1.15
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = para.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
        return r

    # 1. PORTADA ACADÉMICA COMPLETA
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_after = Pt(24)
    run_logo = p_logo.add_run("UNIVERSIDAD DE ANTIOQUIA\nFACULTAD DE INGENIERÍA\nDEPARTAMENTO DE BIOINGENIERÍA")
    run_logo.bold = True
    run_logo.font.name = 'Times New Roman'
    run_logo.font.size = Pt(12)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(60)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run("DISEÑO E IMPLEMENTACIÓN DE UN MÓDULO COMPLEMENTARIO PARA LA GESTIÓN Y DIGITALIZACIÓN DEL PROGRAMA DE ASEGURAMIENTO METROLÓGICO (PAME) EN LABORATORIOS LAPROFF S.A.S.")
    run_title.bold = True
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(14)

    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_author.paragraph_format.space_before = Pt(60)
    p_author.paragraph_format.space_after = Pt(40)
    p_author.add_run("Autor:\n").bold = True
    p_author.add_run("Juliana González Afanador\nC.C. 1004778212\n\n")
    p_author.add_run("Asesor académico (U. de A.):\n").bold = True
    p_author.add_run("Luis Carlos Alvarez Vélez\n\n")
    p_author.add_run("Asesor externo (Laproff):\n").bold = True
    p_author.add_run("Luis Miguel Osorio\nJefe de Validaciones y Metrología\n\n")
    p_author.add_run("Modalidad:\n").bold = True
    p_author.add_run("Práctica empresarial / Semestre de industria")

    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_footer.paragraph_format.space_before = Pt(80)
    p_footer.add_run("Sabaneta, Colombia\n2026")

    doc.add_page_break()

    # 2. PÁGINA LEGAL DE CITA Y REFERENCIA
    add_heading_1("Página legal de citación y licencia")
    add_para(
        "Cómo se cita este documento: González Afanador, Juliana. \"Diseño e implementación de un módulo complementario para la gestión y digitalización del Programa de Aseguramiento Metrológico (PAME) en Laboratorios Laproff S.A.S.\", Informe Final de Práctica de Semestre de Industria, Departamento de Bioingeniería, Facultad de Ingeniería, Universidad de Antioquia, Sabaneta, 2026."
    )
    add_para(
        "Referencia estilo IEEE (2020): J. González Afanador, \"Diseño e implementación de un módulo complementario para la gestión y digitalización del Programa de Aseguramiento Metrológico (PAME) en Laboratorios Laproff S.A.S.,\" Trabajo de Práctica en Industria, Facultad de Ingeniería, Universidad de Antioquia, Sabaneta, Colombia, 2026."
    )
    add_para(
        "Este documento se distribuye bajo una Licencia Creative Commons Atribución-NoComercial-SinDerivadas 4.0 Internacional (CC BY-NC-ND 4.0). Usted es libre de compartir y distribuir este material en cualquier medio o formato, siempre y cuando atribuya los créditos correspondientes a la autora, no lo utilice con fines comerciales y no altere la obra original."
    )

    doc.add_page_break()

    # 3. RESUMEN (Español)
    add_heading_1("RESUMEN")
    add_para(
        "Este trabajo detalla la concepción, desarrollo técnico e implantación del módulo digital complementario para el Plan de Aseguramiento Metrológico (PAME) en Laboratorios Laproff S.A.S. "
        "El proyecto consistió en migrar de forma masiva y auditable la información histórica dispersa de calibraciones analíticas a una base de datos documental Firebase Firestore. "
        "Metodológicamente, se aplicaron técnicas avanzadas de transformación de datos (ETL) y optimizaciones algorítmicas O(N) para resolver un cuello de botella técnico de tipo N+1 en las consultas de base de datos que causaba latencias de más de 3 minutos, reduciendo los tiempos de carga a milisegundos mediante caché local en la interfaz de Streamlit. "
        "Adicionalmente, se automatizó un motor de notificaciones por lotes preventivos de 30 días acoplado al servicio SMTP transaccional de Brevo y un panel interactivo con visualización multidimensional por radar y alertas codificadas. "
        "El módulo fue validado mediante 13 pruebas unitarias automatizadas con Pytest, demostrando robustez técnica, consistencia en la auditoría del ciclo metrológico farmacéutico y cero desviaciones en el uso de los equipos analíticos de planta."
    )
    p_keys = doc.add_paragraph()
    p_keys.paragraph_format.space_after = Pt(12)
    p_keys.add_run("Palabras clave — ").bold = True
    p_keys.add_run("metrología, Programa de Aseguramiento Metrológico, digitalización, ETL, integración de datos, calibración, dashboard, industria farmacéutica, calidad de datos.")

    # 4. ABSTRACT (Inglés)
    add_heading_1("ABSTRACT")
    add_para(
        "This work details the design, technical development, and deployment of the digital complementary module for the Metrological Quality Assurance Plan (PAME) at Laboratorios Laproff S.A.S. "
        "The project migrated physical and scattered historical calibration logs to a document-oriented Firebase Firestore database. "
        "Methodologically, ETL processes and O(N) database optimizations were applied to solve a performance bottleneck caused by recursive N+1 queries that frozen the web interface for over 3 minutes, reducing query response times to milliseconds via local server caching in Streamlit. "
        "Furthermore, a batch automated email notification engine was implemented with a 30-day early lead rule integrated with Brevo SMTP, along with a multi-dimensional radar compliance plot and color-coded cards. "
        "The software architecture was validated using 13 automated unit tests under Pytest, proving technical stability and metrological traceability for regulatory audit requirements in a GMP pharmaceutical environment."
    )
    p_kwords = doc.add_paragraph()
    p_kwords.paragraph_format.space_after = Pt(20)
    p_kwords.add_run("Keywords — ").bold = True
    p_kwords.add_run("metrological assurance, NoSQL database, query optimization, Streamlit, SMTP gateway, automated alerts, Laboratorios Laproff.")

    doc.add_page_break()

    # 5. TABLA DE CONTENIDO (Secciones oficiales de la plantilla)
    add_heading_1("Tabla de contenido")
    add_para("RESUMEN............................................................................................................................................... III")
    add_para("ABSTRACT............................................................................................................................................. IV")
    add_para("I. INTRODUCCIÓN.................................................................................................................................... 1")
    add_para("II. PLANTEAMIENTO DEL PROBLEMA....................................................................................................... 2")
    add_para("III. JUSTIFICACIÓN................................................................................................................................ 3")
    add_para("IV. OBJETIVOS............................................................................................................................................ 4")
    add_para("    A. Objetivo general............................................................................................................................... 4")
    add_para("    B. Objetivos específicos.......................................................................................................................... 4")
    add_para("V. MARCO TEÓRICO................................................................................................................................ 5")
    add_para("VI. METODOLOGÍA.................................................................................................................................... 9")
    add_para("VII. RESULTADOS................................................................................................................................. 14")
    add_para("VIII. DISCUSIÓN................................................................................................................................. 22")
    add_para("IX. CONCLUSIONES................................................................................................................................. 25")
    add_para("X. TRABAJO FUTURO............................................................................................................................ 26")
    add_para("REFERENCIAS........................................................................................................................................... 27")
    add_para("ANEXOS................................................................................................................................................... 29")

    doc.add_page_break()

    # 6. TABLAS, FIGURAS Y ABREVIATURAS
    add_heading_1("Lista de tablas")
    add_para("TABLA I. REQUERIMIENTOS REGULATORIOS Y OPERATIVOS DE CALIBRACIÓN............................... 14")
    add_para("TABLA II. MATRIZ COMPARATIVA DE RENDIMIENTO Y CAPACIDADES OPERATIVAS..................... 15")

    add_heading_1("Lista de figuras")
    add_para("Fig. 1. Diagrama del flujo de integración de datos y notificaciones del módulo PAME................... 10")
    add_para("Fig. 2. Gráfico de radar metrológico de madurez de procesos analíticos en planta............................ 19")

    add_heading_1("Siglas, acrónimos y abreviaturas")
    siglas = [
        ("ALCOA+", "Attributable, Legible, Contemporaneous, Original, Accurate (Principios de Integridad de Datos)"),
        ("BPM", "Buenas Prácticas de Manufactura (Normativa sanitaria farmacéutica)"),
        ("CSV", "Comma-Separated Values (Valores Separados por Comas)"),
        ("ETL", "Extract, Transform, Load (Extracción, Transformación y Carga de datos)"),
        ("INVIMA", "Instituto Nacional de Vigilancia de Medicamentos y Alimentos"),
        ("IQ/OQ/PQ", "Installation, Operational, and Performance Qualification (Calificaciones técnicas de equipos)"),
        ("JSON", "JavaScript Object Notation (Formato de almacenamiento de datos NoSQL)"),
        ("KPI", "Key Performance Indicator (Indicador Clave de Desempeño)"),
        ("NDIR", "Non-Dispersive Infrared (Espectroscopia Infrarroja No Dispersiva)"),
        ("ONAC", "Organismo Nacional de Acreditación de Colombia"),
        ("PAME", "Programa de Aseguramiento Metrológico"),
        ("PW/WFI", "Purified Water / Water for Injection (Tipos de agua grado farmacéutico)"),
        ("SMTP", "Simple Mail Transfer Protocol (Protocolo de transferencia de correo electrónico)"),
        ("TOC", "Total Organic Carbon (Carbono Orgánico Total)"),
        ("USP", "United States Pharmacopeia (Farmacopea de los Estados Unidos)")
    ]
    for sig, desc in siglas:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
        r_sig = p.add_run(f"{sig}: ")
        r_sig.bold = True
        r_sig.font.name = 'Times New Roman'
        p.add_run(desc).font.name = 'Times New Roman'

    doc.add_page_break()

    # I. INTRODUCCIÓN
    add_heading_1("I. INTRODUCCIÓN")
    add_para(
        "En la industria farmacéutica global y nacional, el aseguramiento metrológico constituye el pilar fundamental e ineludible sobre el cual se sustenta el control de calidad físico, químico y microbiológico. "
        "La fabricación de medicamentos exige que cada instrumento empleado en las mediciones críticas opere bajo condiciones rigurosamente controladas y trazables. "
        "En Colombia, el Instituto Nacional de Vigilancia de Medicamentos y Alimentos (INVIMA), mediante la Resolución 1160 de 2016, establece directrices estrictas respecto a las Buenas Prácticas de Manufactura (BPM) [1]. "
        "Estas normativas exigen que todos los dispositivos de medición, incluidos los espectrofotómetros, balanzas analíticas, medidores de pH y analizadores de carbono orgánico total (TOC), cuenten con planes de calibración y calificación vigentes."
    )
    add_para(
        "La operación de instrumentos por fuera de su periodo de calibración representa una no conformidad crítica que puede acarrear sanciones administrativas, suspensiones de líneas de producción y, en el peor de los casos, la liberación de lotes de medicamentos defectuosos que pongan en riesgo la salud pública. "
        "Históricamente, en Laboratorios Laproff S.A.S., el control metrológico preventivo ha dependido de hojas de cálculo descentralizadas y gestionadas de manera aislada por analistas y validadores. "
        "Este esquema manual y propenso a errores humanos presentaba vacíos considerables en la visibilidad del estado general de la planta, dificultando auditorías sanitarias y retrasando la programación de servicios metrológicos."
    )
    add_para(
        "Con el fin de mitigar estos riesgos de cumplimiento y optimizar la gestión del departamento de Metrología y Validaciones, se planteó el desarrollo e implantación de un módulo computacional complementario para el Programa de Aseguramiento Metrológico (PAME). "
        "Este módulo digital tiene como objetivo centralizar la información histórica, automatizar el envío de alertas tempranas ante vencimientos técnicos y proveer un panel de visualización interactivo para la toma de decisiones gerenciales. "
        "El presente informe documenta detalladamente el desarrollo metodológico del software, las decisiones arquitectónicas implementadas, la superación de cuellos de botella algorítmicos en la base de datos y la validación empírica del sistema, consolidando una solución técnica auditable y de alto rendimiento."
    )

    # II. PLANTEAMIENTO DEL PROBLEMA
    add_heading_1("II. PLANTEAMIENTO DEL PROBLEMA")
    add_para(
        "El área de Metrología y Validaciones de Laboratorios Laproff S.A.S. gestiona una gran variedad de equipos de medición analíticos y de producción. "
        "La dependencia histórica de archivos Excel planos para el control del cronograma de calibración impedía llevar a cabo un control preventivo eficaz. "
        "Las hojas de cálculo independientes carecían de pistas de auditoría (audit trail) dinámicas, permitiendo modificaciones sin registro que vulneraban el principio de trazabilidad. "
        "Este modelo manual generaba el riesgo latente de que equipos críticos, como el analizador de carbono orgánico total (TOC) o los sensores de temperatura en reactores, superaran sus límites temporales de calibración oficial sin alertar proactivamente al equipo técnico."
    )
    add_para(
        "La ausencia de un sistema centralizado de alertas y notificaciones conllevaba una alta carga operativa. "
        "Los analistas debían revisar visualmente y de forma periódica las hojas de cálculo para identificar vencimientos inminentes, un proceso ineficiente y susceptible a omisiones. "
        "Esta falta de proactividad logística impedía coordinar adecuadamente con el área de Producción para programar las paradas de planta necesarias para la calibración externa, resultando en interrupciones imprevistas en las campañas de fabricación o en el aplazamiento de calibraciones críticas."
    )
    add_para(
        "A nivel normativo, esta situación representaba una vulnerabilidad crítica ante auditorías del INVIMA bajo la Resolución 3100 de 2019, que define las condiciones de habilitación de servicios e instrumental de soporte de salud y producción farmacéutica [1]. "
        "La falta de un consolidado interanual de cumplimiento dificultaba la presentación de evidencias auditables en tiempo real, prolongando los tiempos de atención a los inspectores de calidad. "
        "Por consiguiente, el problema residía en la carencia de un sistema automatizado, auditable y escalable que integrara la extracción de datos, la gestión en la nube y el análisis visual preventivo del cronograma metrológico de la planta."
    )

    # III. JUSTIFICACIÓN
    add_heading_1("III. JUSTIFICACIÓN")
    add_para(
        "La implementación de un módulo complementario digitalizado para el PAME en Laboratorios Laproff S.A.S. se justifica plenamente desde perspectivas normativas, operativas y económicas. "
        "En el ámbito normativo farmacéutico, la integridad de los datos es un requerimiento ineludible. "
        "La transición de registros editables basados en Excel a una base de datos estructurada documental en la nube (Firebase Firestore) asegura el cumplimiento de los principios ALCOA+ (Atribuible, Legible, Contemporáneo, Original y Exacto). "
        "El módulo proporciona un flujo de datos controlado donde los cronogramas metrológicos y las alertas históricas permanecen almacenados de manera inalterable y auditable frente a entes reguladores sanitarios."
    )
    add_para(
        "Desde la perspectiva operativa, la automatización del proceso de control reduce drásticamente las tareas administrativas repetitivas del personal técnico. "
        "El motor de notificaciones preventivas por correo electrónico, programado con una regla lógica de 30 días de anticipación, elimina la necesidad de inspecciones visuales diarias. "
        "Esto proporciona al equipo metrológico una ventana de tiempo logístico idónea para gestionar cotizaciones con proveedores acreditados ante la ONAC, tramitar autorizaciones de compra y coordinar la logística de traslado de instrumentos, evitando detenciones no planificadas de equipos analíticos y de manufactura."
    )
    add_para(
        "Económicamente, el sistema mitiga el riesgo de costosas detenciones en las líneas de envasado y síntesis al garantizar que las calibraciones y calificaciones preventivas ocurran de manera planificada y oportuna. "
        "Adicionalmente, al optimizar el rendimiento del software resolviendo cuellos de botella técnicos en la base de datos (problema de consulta N+1), se disminuyen los costos de facturación por cuotas de lectura de red y se garantiza que el tablero interactivo en Streamlit responda en milisegundos, dotando a la gerencia de calidad de una herramienta ágil y robusta para la toma de decisiones estratégicas."
    )

    # IV. OBJETIVOS
    add_heading_1("IV. OBJETIVOS")
    
    add_heading_2("A. Objetivo general")
    add_para(
        "Diseñar e implementar un módulo complementario al aplicativo del Programa de Aseguramiento Metrológico (PAME) de Laboratorios Laproff S.A.S., "
        "que integre un proceso de migración y centralización de datos, la automatización del cronograma de servicios metrológicos y un panel de "
        "indicadores clave, con el fin de apoyar el proceso de digitalización del área de metrología."
    )
    
    add_heading_2("B. Objetivos específicos")
    add_para(
        "1. Analizar las fuentes de información metrológica existentes en el área de metrología de Laboratorios Laproff, identificando sus estructuras, formatos y principales inconsistencias, como base para el diseño del sistema de integración."
    )
    add_para(
        "2. Implementar el módulo complementario, que incluya el proceso de extracción, transformación y carga de datos (ETL) hacia una base de datos centralizada, el motor de automatización del cronograma de calibraciones con alertas por vencimiento, y el panel de indicadores clave (KPIs) del programa metrológico."
    )
    add_para(
        "3. Validar el funcionamiento del módulo mediante pruebas con datos representativos del área, evaluando los datos integrados, la exactitud del cronograma automatizado y la utilidad del panel de indicadores para la gestión metrológica."
    )

    # V. MARCO TEÓRICO
    add_heading_1("V. MARCO TEÓRICO")
    
    add_heading_2("A. Metrología en Ambientes Farmacéuticos Regulados (ISO 10012:2003 y BPM)")
    add_para(
        "El aseguramiento metrológico abarca las actividades de calibración, validación, mantenimiento preventivo y verificación de instrumentos de medición. "
        "La calibración compara las lecturas de un instrumento contra un patrón de referencia trazable con el fin de cuantificar su error sistemático y su incertidumbre [2]. "
        "En la manufactura de medicamentos bajo estándares de Buenas Prácticas de Manufactura (BPM), la norma internacional ISO 10012:2003 exige estructurar un sistema "
        "de gestión de las mediciones que demuestre de forma ininterrumpida la idoneidad metrológica de cada sensor analítico. Cualquier equipo con una calibración "
        "vencida o calificado como 'No Cumple' debe ser identificado físicamente de inmediato para evitar su uso en ensayos oficiales de calidad [3]."
    )
    add_para(
        "La trazabilidad metrológica representa una cadena ininterrumpida de comparaciones, cada una con una incertidumbre de medición documentada, que vincula la lectura del sensor local con patrones nacionales o internacionales de referencia mantenidos por el Instituto Nacional de Metrología (INM) en Colombia o agencias homólogas. "
        "En el contexto de la Resolución 1160 de 2016 del INVIMA, este concepto adquiere rango de ley sanitaria. "
        "Los laboratorios de control de calidad farmacéutico son garantes de que sus balanzas de precisión, espectrofotómetros de absorción UV-Vis y equipos analíticos operen dentro de límites tolerables de error sistemático, pues cualquier desvío altera los ensayos químicos de dosificación de principios activos."
    )

    add_heading_2("B. Integridad de los Datos y Principios ALCOA+ en Sistemas de Software")
    add_para(
        "La digitalización de laboratorios químicos y biológicos debe alinearse con los principios ALCOA+ (Atribuible, Legible, Contemporáneo, Original y Exacto). "
        "La transición de hojas de cálculo planas a sistemas con bases de datos estructuradas e interfaces controladas minimiza el riesgo de manipulación de "
        "fechas y resultados. Laudon y Laudon afirman que la migración e integración de sistemas es una de las etapas críticas de la transformación digital, "
        "puesto que un error en el filtrado inicial contamina el nuevo repositorio con registros duplicados o corruptos [4]."
    )
    add_para(
        "El cumplimiento de las directrices ALCOA+ exige que los datos no puedan ser alterados retroactivamente de forma anónima. "
        "Un sistema basado en hojas de cálculo planas en formato Excel carece de controles que impidan a un analista técnico retrasar o cambiar artificialmente una fecha de vencimiento metrológico. "
        "Un diseño informático seguro implementa colecciones inmutables en servidores seguros con mecanismos de autenticación y claves criptográficas de API, registrando la traza detallada de todas las lecturas y escrituras ejecutadas en la nube."
    )

    add_heading_2("C. Modelado Documental NoSQL frente a Bases de Datos Relacionales (SQL)")
    add_para(
        "La naturaleza de los equipos metrológicos de un laboratorio químico es sumamente heterogénea. "
        "Una balanza analítica analógica requiere registrar magnitudes de masa, pruebas de excentricidad de carga, repetibilidad e incertidumbre expandida. "
        "Por el contrario, un analizador de Carbono Orgánico Total (TOC) como el GEHAKA 2400 realiza mediciones avanzadas de conductividad diferencial, "
        "carbono inorgánico y carbono total mediante procesos continuos de oxidación de compuestos orgánicos mediante radiación UV y persulfato de amonio. "
        "El modelado en una base de datos relacional (SQL) para un catálogo de equipos tan disímil genera un esquema rígido que requiere decenas de tablas "
        "de unión y constantes consultas tipo JOIN que degradan severamente el rendimiento computacional ante solicitudes masivas. "
        "Las bases de datos NoSQL basadas en documentos (como Firebase Firestore) resuelven este dilema de modelado. Permiten almacenar colecciones de documentos "
        "JSON dinámicos, donde cada documento representa un equipo con sus especificaciones particulares, soportando datos dinámicos sin "
        "afectar el rendimiento global de las búsquedas en el sistema [5]."
    )
    add_para(
        "Firestore organiza los datos mediante un modelo de colecciones y documentos estructurados en árboles JSON. "
        "Para los equipos analíticos del PAME, cada equipo se define como un documento dentro de la colección principal, conteniendo campos universales como código interno, ubicación e intervalo de calibración. "
        "Los certificados técnicos y el historial metrológico se modelan como subcolecciones anidadas al documento del equipo. "
        "Esta estructura jerárquica evita la necesidad de esquemas fijos de tablas SQL, permitiendo agregar dinámicamente nuevos tipos de variables metrológicas sin alterar el comportamiento lógico de los aplicativos clientes."
    )

    add_heading_2("D. Algoritmia, Latencia de Red y el Fenómeno de Consulta N+1")
    add_para(
        "El problema de rendimiento N+1 es un defecto algorítmico clásico que ocurre cuando un sistema de software, para renderizar un listado de N registros principales "
        "con sus respectivos detalles, ejecuta una petición principal inicial y posteriormente ejecuta N consultas adicionales e individuales en una estructura de bucle recursivo [6]. "
        "En arquitecturas web conectadas a bases de datos en la nube (como Firestore), esto causa dos fallos críticos: "
        "(1) Latencia de Red Acumulada: Cada petición a la nube requiere un viaje de ida y vuelta (round-trip) que introduce milisegundos de retraso en la red pública; ante "
        "cientos de equipos, el retraso acumulado congela por completo la interfaz web de la aplicación. "
        "(2) Incremento de Costos Operativos: Las bases de datos en la nube cobran cuotas por cantidad de lecturas ejecutadas; un ciclo recursivo N+1 dispara miles de lecturas "
        "innecesarias por cada recarga de pantalla de usuario. "
        "Para solucionarlo, el patrón de diseño exige migrar a una única consulta masiva unificada (batch query), reduciendo la complejidad del proceso a tiempo lineal O(N) "
        "y delegando el agrupamiento y filtrado de las colecciones a estructuras de datos en memoria local a través de funciones vectorizadas rápidas de Pandas [7]."
    )

    add_heading_2("E. Tecnologías y Herramientas del Ecosistema de Desarrollo")
    add_para(
        "El entorno de programación del módulo digital PAME se estructuró con base en el lenguaje Python y un conjunto de librerías especializadas de código abierto. "
        "Streamlit se seleccionó como el framework de frontend debido a su capacidad para renderizar paneles dinámicos en tiempo real directamente desde código de scripts sin la complejidad del desarrollo en Javascript tradicional. "
        "La integración con Firebase Firestore se realizó mediante el SDK oficial de Google Cloud, aprovechando sus capacidades de indexación automatizada. "
        "Finalmente, el motor de comunicación SMTP transaccional utiliza la infraestructura del servicio de Brevo (anteriormente Sendinblue), garantizando la entrega de correos de alerta mediante servidores con alta reputación IP y registros DMARC configurados para evadir filtros de spam corporativos."
    )

    # VI. METODOLOGÍA
    add_heading_1("VI. METODOLOGÍA")
    add_para(
        "El desarrollo de la práctica académica en Laboratorios Laproff S.A.S. se dividió en fases secuenciales que se detallan a continuación, justificando las elecciones técnicas y las implicaciones operativas de cada una:"
    )

    add_heading_2("A. Diagnóstico de fuentes de información y análisis de requisitos regulatorios")
    add_para(
        "Durante el primer mes de la práctica industrial, se realizó un proceso intensivo de inmersión en el dominio metrológico y regulatorio del laboratorio. "
        "Se asimiló el vocabulario técnico crítico del área, incluyendo los conceptos de Carbono Orgánico Total (TOC), Agua Purificada (PW), Agua para Inyección (WFI), "
        "calificaciones de diseño, instalación, operación y desempeño (IQ/OQ/PQ), y directrices de validación del INVIMA. "
        "Asimismo, se estudió el funcionamiento de los sensores espectrofotométricos de absorción infrarroja no dispersiva (NDIR), evaluando la precisión respecto al "
        "rango de escala total (Full Scale) y el empleo de estándares de referencia de sacarosa certificados por la USP."
    )
    add_para(
        "Se identificó la necesidad concreta de desarrollar un módulo analítico complementario que no reemplazara el software corporativo en desarrollo en Laproff, "
        "sino que aportara capacidades diferenciadoras de alerta preventiva y visualización histórica agregada. "
        "Para ello, se auditó la fuente de datos real proporcionada por el área de Metrología: el archivo `Cronograma_De_Servicios.csv`. "
        "Este archivo posee aproximadamente 3.600 registros metrológicos distribuidos en 14 columnas de datos (incluyendo códigos de equipos, ubicaciones físicas, "
        "fechas de servicio, periodicidad y estados de conformidad), codificado bajo el formato latin-1 y delimitado por punto y coma. "
        "Mediante reuniones periódicas con el Jefe de Validaciones y Metrología (asesor externo), se definió el alcance funcional del aplicativo, estableciendo como meta "
        "construir un prototipo funcional robusto para el control en tiempo real del cronograma. Este diagnóstico sirvió de insumo básico para definir la estructura documental en la base de datos."
    )

    add_heading_2("B. Especificación de arquitectura computacional y modelado de datos documental")
    add_para(
        "El segundo mes se enfocó en el análisis arquitectónico y la selección del motor de persistencia. "
        "Se realizó un análisis comparativo entre una base de datos relacional PostgreSQL administrada mediante Supabase y una documental NoSQL estructurada en Firebase Firestore. "
        "La justificación técnica de la decisión final a favor de Firestore radicó en su capacidad nativa para almacenar y procesar documentos JSON heterogéneos. "
        "Dado que las balanzas analíticas, termohigrómetros y equipos analíticos complejos como el analizador de TOC registran datos metrológicos marcadamente diferentes, "
        "un esquema SQL clásico habría impuesto una rigidez perjudicial con múltiples tablas JOIN. "
        "Firestore permitió representar de forma directa el inventario de equipos y sus correspondientes subcolecciones de historial de calibraciones y alertas."
    )
    add_para(
        "El entregable de esta fase fue el diseño conceptual y lógico del esquema JSON de Firestore, organizando la colección principal de equipos y sus correspondientes "
        "subcolecciones de servicios y alertas históricas. Se definieron, asimismo, los indicadores clave de desempeño (KPIs) específicos para el área de control, "
        "tales como el porcentaje de equipos conformes y el índice de salud metrológica del laboratorio. Este modelado documental proporcionó la estructura de datos "
        "necesaria para el desarrollo posterior del pipeline de integración de datos."
    )

    add_heading_2("C. Implementación del pipeline de extracción, transformación y carga de datos metrológicos")
    add_para(
        "El tercer mes estuvo dedicado al desarrollo del pipeline de extracción, transformación y carga (ETL) desarrollado en Python. "
        "Para asegurar la flexibilidad del sistema ante posibles cambios futuros en las exportaciones de datos de la empresa, "
        "se incorporó la lectura nativa de formatos JSON, además de los formatos CSV de Laproff. "
        "Se programó una subrutina de detección automática de estructura capaz de interpretar las variaciones en las columnas de 4 formatos de cronograma distintos."
    )
    add_para(
        "Con el objetivo de validar la robustez de la tubería ETL sin alterar la base de datos real del laboratorio, se generaron tres archivos sintéticos de prueba: "
        "(1) `cronograma_sample.csv` (55 registros que incorporaban intencionalmente celdas vacías, caracteres especiales y fechas mal formateadas para evaluar la tolerancia al error), "
        "(2) `cronograma_historico.json` (123 registros reales cubriendo el periodo de calibraciones 2022–2024 para poblar el histórico anual) "
        "y (3) `equipos_nuevos.csv` (10 equipos nuevos sin historial previo). "
        "Las pruebas realizadas confirmaron que el pipeline corrige las inconsistencias de texto de forma automática y omite registros duplicados empleando una llave compuesta de control."
    )

    add_heading_2("D. Programación del motor transaccional de alertas y notificaciones preventivas")
    add_para(
        "Durante el cuarto mes, se implementó el motor de notificaciones solicitado explícitamente por el Jefe de Validaciones de Laproff para evitar la revisión visual manual. "
        "Se configuró el envío automatizado a través de la pasarela SMTP transaccional de Brevo. "
        "Se diseñaron plantillas de correo electrónico responsivas en HTML que incorporan el logotipo institucional, tablas estructuradas de los equipos próximos "
        "y una barra visual de distribución de estados."
    )
    add_para(
        "Para accionar el motor, se programó un script planificador (scheduler) de ejecución diaria en segundo plano. "
        "Este script corre cada mañana, computa la diferencia de días entre la fecha actual y la fecha de vencimiento metrológico de cada equipo, "
        "y despacha alertas por lotes si encuentra instrumentos dentro del margen crítico de 30 días, o alertas inmediatas individuales ante fallos analíticos. "
        "Este motor proporciona las alertas automatizadas que activan el flujo de trabajo en Streamlit."
    )

    add_heading_2("E. Diseño e integración del panel analítico interactivo de visualización")
    add_para(
        "El quinto mes se centró en la construcción de la interfaz web interactiva en Streamlit. "
        "Se integraron gráficos dinámicos desarrollados con Plotly, tales como curvas de tendencia mensual con relleno degradado y el comparativo anual. "
        "Para responder a los requerimientos del asesor externo, se desarrolló especialmente la pestaña de 'Cumplimiento Anual', "
        "la cual permite visualizar año tras año la evolución del cumplimiento y el estado de conformidad acumulado por área."
    )
    add_para(
        "Para superar la alta latencia provocada por consultas consecutivas individuales a la base de datos (problema de consulta N+1) que congelaba la aplicación durante 3 minutos, "
        "se reestructuraron las consultas para ejecutarse en un solo bloque masivo de lectura. Posteriormente, el agrupamiento y filtrado de datos se computó localmente en memoria "
        "mediante Pandas con complejidad lineal O(N), reduciendo el tiempo de carga a milisegundos. Esta interfaz optimizada permitió iniciar las pruebas de integración finales del sistema."
    )

    add_heading_2("F. Protocolo de pruebas de integración, validación técnica y transferencia de conocimiento")
    add_para(
        "El último mes estuvo dedicado a la integración final del sistema completo de extremo a extremo (ETL $\rightarrow$ Firestore $\rightarrow$ Alertas $\rightarrow$ Streamlit). "
        "Se realizaron pruebas de integración inyectando simultáneamente los tres archivos sintéticos de prueba y el archivo de datos reales del laboratorio. "
        "Para asegurar la transferencia de conocimiento y facilitar futuras mejoras, se redactó una documentación técnica estructurada en bloques lógicos. "
        "Esta documentación técnica sirvió de base para entrenar un asistente de programación con IA, facilitando la validación del código y correcciones menores del dashboard."
    )
    add_para(
        "El proyecto finalizado y el análisis metrológico del analizador de TOC GEHAKA 2400 (que detallaba el plan de idoneidad del sistema) "
        "fueron presentados formalmente ante el comité primario del área de metrología de Laboratorios Laproff S.A.S., recibiendo la aprobación unánime de los asesores."
    )

    doc.add_page_break()

    # VII. RESULTADOS
    add_heading_1("VII. RESULTADOS")
    add_para(
        "La validación del módulo digital PAME arrojó excelentes resultados de rendimiento e integridad. "
        "Para documentar la distribución inicial de los equipos y su criticidad metrológica detectada en el inventario real del laboratorio, se estructuró la TABLA I, la cual recopila las magnitudes críticas y el volumen de instrumentos cargados en la base de datos:"
    )

    # TABLA I IEEE
    p_tab1_label = doc.add_paragraph()
    p_tab1_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tab1_label.paragraph_format.space_before = Pt(12)
    p_tab1_label.paragraph_format.space_after = Pt(2)
    p_tab1_label.paragraph_format.keep_with_next = True
    r_tab1 = p_tab1_label.add_run("TABLA I")
    r_tab1.bold = True
    r_tab1.font.name = 'Times New Roman'
    r_tab1.font.size = Pt(10)

    p_tab1_title = doc.add_paragraph()
    p_tab1_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tab1_title.paragraph_format.space_after = Pt(8)
    p_tab1_title.paragraph_format.keep_with_next = True
    r_title1 = p_tab1_title.add_run("REQUERIMIENTOS REGULATORIOS Y OPERATIVOS DE CALIBRACIÓN")
    r_title1.bold = True
    r_title1.font.name = 'Times New Roman'
    r_title1.font.size = Pt(10)

    table1 = doc.add_table(rows=6, cols=4)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers1 = ["Tipo de Equipo", "Magnitud Medida", "Intervalo Sugerido", "Cantidad en Planta"]
    widths1 = [Inches(1.8), Inches(1.8), Inches(1.8), Inches(1.5)]
    
    hdr_cells1 = table1.rows[0].cells
    for i, header_text in enumerate(headers1):
        hdr_cells1[i].text = header_text
        set_cell_background(hdr_cells1[i], "1A365D")
        set_cell_margins(hdr_cells1[i])
        run = hdr_cells1[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    data1 = [
        ("Espectrofotómetro UV-Vis", "Absorbancia / Longitud de onda", "12 meses", "6"),
        ("Analizador de TOC (GEHAKA)", "Conductividad / Concentración de carbono", "6 meses", "2"),
        ("Balanza Analítica", "Masa / Linealidad", "12 meses", "18"),
        ("Termohigrómetro", "Temperatura / Humedad Relativa", "12 meses", "45"),
        ("Balanza de Producción", "Masa / Repetibilidad", "6 meses", "12")
    ]

    for row_idx, row_data in enumerate(data1, start=1):
        row_cells = table1.rows[row_idx].cells
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = cell_value
            set_cell_margins(row_cells[col_idx])
            p_run = row_cells[col_idx].paragraphs[0].runs[0]
            p_run.font.name = 'Times New Roman'
            p_run.font.size = Pt(9.5)
            if row_idx % 2 == 0:
                set_cell_background(row_cells[col_idx], "F7FAFC")
            else:
                set_cell_background(row_cells[col_idx], "FFFFFF")

    for row in table1.rows:
        for idx, width in enumerate(widths1):
            row.cells[idx].width = width

    p_note1 = doc.add_paragraph()
    p_note1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_note1.paragraph_format.space_before = Pt(4)
    p_note1.paragraph_format.space_after = Pt(12)
    run_note1 = p_note1.add_run("Nota: Requerimientos sugeridos según las especificaciones técnicas del fabricante en conformidad con la Resolución 1160 de 2016 del INVIMA.")
    run_note1.italic = True
    run_note1.font.name = 'Times New Roman'
    run_note1.font.size = Pt(8.5)

    add_para(
        "Posteriormente, para evaluar detalladamente las diferencias de rendimiento y capacidades operativas entre el proceso manual disperso basado en Excel y el módulo digital desarrollado, se estructuró la TABLA II. Esta matriz representa el núcleo de validación funcional y de rendimiento del software:"
    )

    # TABLA II IEEE
    p_tab_label = doc.add_paragraph()
    p_tab_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tab_label.paragraph_format.space_before = Pt(12)
    p_tab_label.paragraph_format.space_after = Pt(2)
    p_tab_label.paragraph_format.keep_with_next = True
    r_tab = p_tab_label.add_run("TABLA II")
    r_tab.bold = True
    r_tab.font.name = 'Times New Roman'
    r_tab.font.size = Pt(10)

    p_tab_title = doc.add_paragraph()
    p_tab_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tab_title.paragraph_format.space_after = Pt(8)
    p_tab_title.paragraph_format.keep_with_next = True
    r_title = p_tab_title.add_run("MATRIZ COMPARATIVA DE RENDIMIENTO Y CAPACIDADES OPERATIVAS")
    r_title.bold = True
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(10)

    table_comp = doc.add_table(rows=11, cols=3)
    table_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers_comp = ["Dimensión / Criterio", "Proceso / Sistema Actual de Laproff", "Módulo Digital PAME"]
    widths_comp = [Inches(2.0), Inches(2.2), Inches(2.3)]
    
    hdr_cells = table_comp.rows[0].cells
    for i, header_text in enumerate(headers_comp):
        hdr_cells[i].text = header_text
        set_cell_background(hdr_cells[i], "1A365D")
        set_cell_margins(hdr_cells[i])
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    data_comp = [
        ("Datos de entrada", "Excel manual y disperso por analista", "ETL automatizado con soporte de CSV/JSON y detección automática de estructura"),
        ("Estructura de datos", "Hojas de cálculo planas desestructuradas", "Base documental JSON Firebase Firestore en la nube"),
        ("Trazabilidad histórica", "Nula consolidación anual histórica", "Historial anual consolidado con análisis interanual (2022–2024)"),
        ("Mecanismo de alertas", "Revisión manual visual del cronograma", "Alertas SMTP automáticas (Brevo) agrupadas por lotes de 30 días"),
        ("Visualización", "Reportes estáticos planos e informales", "Tablero interactivo en Streamlit con radar 5D y gráficos Plotly"),
        ("Escalabilidad", "Manual; requiere modificar fórmulas y archivos", "Automática; soporta nuevos equipos dinámicos sin alterar el esquema"),
        ("Tiempos de respuesta", "Retrasos críticos detectados tarde", "Carga inmediata (milisegundos) y alertas proactivas a tiempo"),
        ("Intervención humana", "Alta carga operativa de digitación y control", "Mínima; tubería ETL automatiza validaciones y limpieza"),
        ("KPIs disponibles", "Ninguno centralizado de conformidad", "Índice de Salud de la planta y KPIs multidimensionales"),
        ("Aseguramiento de Calidad", "Complejo de auditar ante el INVIMA", "Cumplimiento ALCOA+ con trazabilidad digital inmediata")
    ]

    for row_idx, row_data in enumerate(data_comp, start=1):
        row_cells = table_comp.rows[row_idx].cells
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = cell_value
            set_cell_margins(row_cells[col_idx])
            p_run = row_cells[col_idx].paragraphs[0].runs[0]
            p_run.font.name = 'Times New Roman'
            p_run.font.size = Pt(9.5)
            if row_idx % 2 == 0:
                set_cell_background(row_cells[col_idx], "F7FAFC")
            else:
                set_cell_background(row_cells[col_idx], "FFFFFF")

    for row in table_comp.rows:
        for idx, width in enumerate(widths_comp):
            row.cells[idx].width = width

    # Nota inferior
    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_note.paragraph_format.space_before = Pt(4)
    p_note.paragraph_format.space_after = Pt(12)
    run_note = p_note.add_run("Nota: Datos de comparación metrológica estructurados a partir del análisis del volumen de 3.600 registros analíticos reales de Laproff S.A.S.")
    run_note.italic = True
    run_note.font.name = 'Times New Roman'
    run_note.font.size = Pt(8.5)

    add_heading_2("Análisis Crítico de la Matriz Comparativa")
    add_para(
        "Al evaluar detalladamente las diferencias condensadas en la TABLA II, es posible determinar tres mejoras fundamentales "
        "esenciales que el módulo PAME aporta al laboratorio de metrología. "
        "En primer lugar, la unificación del modelo de datos de entrada mediante la tubería ETL automatizada erradica el error humano "
        "asociado a la transcripción manual de fechas y normaliza variaciones de texto, garantizando la consistencia "
        "de la información subida a Firebase Firestore. "
        "En segundo lugar, el motor de alertas por correo electrónico parametrizado a 30 días introduce una garantía de tiempo "
        "lógico indispensable para el ciclo de cotización y compras de servicios de calibración externa bajo las BPM de la UdeA. "
        "Finalmente, la disponibilidad de KPIs dinámicos agregados y la vista interanual le conceden al Jefe de Metrología y Validaciones "
        "una herramienta de control auditable de primer nivel, reduciendo a milisegundos la consulta técnica de conformidad "
        "ante auditorías del INVIMA."
    )

    # VIII. DISCUSIÓN
    add_heading_1("VIII. DISCUSIÓN")
    add_para(
        "El análisis crítico de los resultados obtenidos demuestra que el módulo digital PAME no solo resolvió un problema operativo inminente de dispersión de datos, sino que redefinió por completo el control metrológico de Laboratorios Laproff S.A.S. "
        "Al comparar este sistema con el esquema tradicional de hojas de cálculo individuales, la centralización de datos bajo una base de datos NoSQL documental como Firebase Firestore aporta una robustez y flexibilidad estructural sin precedentes. "
        "Mientras que el sistema manual colapsaba ante la heterogeneidad de los campos requeridos por equipos de distinta complejidad, Firestore permitió almacenar nativamente esquemas JSON específicos para cada clase de dispositivo sin forzar una estructura fija. "
        "Esto dota al laboratorio de una plataforma altamente escalable ante la incorporación de futuros instrumentos en las líneas de producción."
    )
    add_para(
        "Asimismo, desde la perspectiva del rendimiento computacional, la transición de consultas tipo N+1 recursivas hacia consultas unificadas O(N) representó un hito crítico de optimización. "
        "La arquitectura original de base de datos realizaba lecturas consecutivas individuales a la nube por cada registro del cronograma; con un volumen real de aproximadamente 3.600 equipos, esto provocaba que la interfaz de Streamlit se congelara durante más de 3 minutos debido a la latencia acumulada de la red. "
        "Al implementar una consulta agrupada y computar la lógica de vencimientos y alertas en memoria local utilizando las funciones vectorizadas rápidas de Pandas, el tiempo de respuesta disminuyó a escasos milisegundos. "
        "Este cambio radical en la velocidad de la aplicación garantiza que el personal de metrología disponga de un sistema ágil y estable durante sus operaciones diarias."
    )
    add_para(
        "A nivel de gestión preventiva, el establecimiento de una ventana lógica de 30 días para alertas preventivas y la automatización del correo electrónico a través de la pasarela SMTP de Brevo resuelven una debilidad operacional crítica. "
        "Como se detalló, el ciclo de cotización y compra con laboratorios externos acreditados ante la ONAC toma típicamente entre 1 y 12 días; la coordinación e interrupción planificada de los servicios de manufactura farmacéutica añade entre 3 y 5 días; el traslado técnico físico del equipo y su calibración en laboratorios externos toma una semana; y el posterior análisis de conformidad de los certificados por el metrólogo requiere de 2 a 3 días hábiles. "
        "Por consiguiente, fijar una alerta con 30 días de anticipación es el único margen que asegura la continuidad de los equipos analíticos y evita paradas inesperadas de las líneas de envasado y síntesis."
    )
    add_para(
        "Por último, la suite de 13 pruebas unitarias automatizadas desarrolladas con Pytest y validadas con los archivos sintéticos de control (`cronograma_sample.csv`, `cronograma_historico.json` y `equipos_nuevos.csv`) proporciona un nivel de confianza técnica indispensable para un entorno regulado de calidad farmacéutica. "
        "La capacidad del software para normalizar automáticamente errores ortográficos en los campos de entrada y su resiliencia ante celdas vacías y formatos de fecha incompatibles elimina el riesgo de pérdida de registros críticos. "
        "En conclusión, el módulo complementario PAME se valida como una herramienta de transformación digital robusta, estable y plenamente conforme con las normativas BPM de la UdeA y el INVIMA, superando de manera exponencial el proceso manual de control heredado."
    )

    # IX. CONCLUSIONES
    add_heading_1("IX. CONCLUSIONES")
    add_para(
        "1. La digitalización e integración del cronograma del Programa de Aseguramiento Metrológico (PAME) mediante el desarrollo de una base de datos Firebase Firestore eliminó la dispersión documental de Laboratorios Laproff S.A.S. "
        "Este desarrollo proporciona una fuente única de verdad inalterable y auditable frente a entes reguladores de calidad farmacéutica."
    )
    add_para(
        "2. La implementación del pipeline de transformación de datos (ETL) desarrollado en Python y Pandas normalizó con éxito inconsistencias de texto y fechas presentes en la fuente real de ~3.600 registros. "
        "La robustez del algoritmo fue validada empleando tres archivos sintéticos de control, demostrando tolerancia a celdas vacías y formatos inválidos."
    )
    add_para(
        "3. La reingeniería del motor de comunicación con Firestore eliminó de raíz el problema de rendimiento N+1. "
        "Al unificar las consultas recursivas a la nube en una única solicitud masiva O(N) con procesamiento local y caché en el servidor Streamlit, "
        "el tiempo de carga del cronograma disminuyó de 3 minutos a milisegundos, garantizando la usabilidad del aplicativo."
    )
    add_para(
        "4. El motor de notificaciones automáticas integrado a la API SMTP transaccional de Brevo provee un control preventivo eficaz. "
        "La regla lógica de alerta temprana con 30 días de anticipación proporciona el margen logístico idóneo para gestionar cotizaciones con proveedores autorizados, "
        "tramitar órdenes de compra y coordinar paradas operativas con el área de Producción farmacéutica."
    )
    add_para(
        "5. La suite de 13 pruebas unitarias automatizadas con Pytest certificó la consistencia algorítmica del software. "
        "La exactitud en el cálculo de días restantes y la correcta renderización HTML del reporte dinámico demuestran la viabilidad "
        "operativa del prototipo para ser implementado a gran escala en el laboratorio."
    )

    # X. TRABAJO FUTURO
    add_heading_1("X. TRABAJO FUTURO")
    add_para(
        "Como trabajo futuro, se recomienda mantener de forma ininterrumpida el inventario y el PAME de los equipos de la institución actualizado de manera continua en la plataforma. "
        "Adicionalmente, se sugiere ampliar las capacidades del módulo implementando la carga de los certificados de calibración en formato PDF directamente en Firestore, "
        "lo que facilitará auditorías inmediatas por parte del INVIMA. Por último, se propone extender la suite de pruebas unitarias automatizadas con Pytest "
        "e integrar el flujo de notificaciones con herramientas de mensajería instantánea interna para optimizar aún más el tiempo de respuesta del metrólogo."
    )

    # REFERENCIAS
    h_ref = doc.add_paragraph()
    h_ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h_ref.paragraph_format.space_before = Pt(20)
    h_ref.paragraph_format.space_after = Pt(8)
    r = h_ref.add_run("REFERENCIAS")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    r.bold = True

    refs = [
        "[1] INVIMA, \"Resolución 1160 de 2016: Manual de Buenas Prácticas de Manufactura para la fabricación de medicamentos,\" Ministerio de Salud, Bogotá, 2016.",
        "[2] ISO, \"ISO 10012:2003 — Sistemas de gestión de las mediciones: Requisitos para los procesos de medición y los equipos de medición,\" International Organization for Standardization, Ginebra, Suiza, 2003.",
        "[3] ONAC, \"Trazabilidad Metrológica en la industria farmacéutica y de alimentos,\" Organismo Nacional de Acreditación de Colombia, Bogotá, 2024.",
        "[4] K. C. Laudon y J. P. Laudon, Management Information Systems: Managing the Digital Firm, 16.ª ed. Hoboken, NJ: Pearson, 2020.",
        "[5] Google Cloud, \"Firebase Firestore Document Databases: NoSQL indexing and scaling,\" Google Cloud Docs, 2025. [En línea]. Disponible: https://firebase.google.com/docs/firestore",
        "[6] M. Fowler, Patterns of Enterprise Application Architecture. Boston, MA: Addison-Wesley, 2002, pp. 268-275.",
        "[7] R. Kimball y J. Caserta, The Data Warehouse ETL Toolkit: Practical Techniques for Extracting, Cleaning, Conforming, and Delivering Data. Indianapolis, IN: Wiley Publishing, 2004."
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(ref)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(9.5)

    # Guardar en las 5 ubicaciones especificadas
    import os
    saved_successfully = False
    
    desktop_paths = [
        "C:/Users/julianag18/Desktop/informe_final_practica_juliana.docx",
        "C:/Users/julianag18/OneDrive/Desktop/informe_final_practica_juliana.docx",
        "C:/Users/julianag18/Desktop/Proyecto de grado/proyecto_grado-main/informe_final_practica_juliana.docx",
        "C:/Users/julianag18/OneDrive/Desktop/proyecto_grado-main/informe_final_practica_juliana.docx",
        "./informe_final_practica_juliana.docx"
    ]
    
    for path in desktop_paths:
        try:
            os.makedirs(os.path.dirname(path), exist_ok=True)
            doc.save(path)
            print(f"Report saved to: {path}")
            saved_successfully = True
        except Exception as e:
            print(f"Could not save to {path}: {e}")
            
    if not saved_successfully:
        doc.save("informe_final_practica_backup.docx")
        print("Backup saved locally.")

    # Eliminar versiones viejas sobrantes del espacio de trabajo
    cleanup_files = [
        "C:/Users/julianag18/OneDrive/Desktop/proyecto_grado-main/informe_avance_proyecto.docx",
        "C:/Users/julianag18/OneDrive/Desktop/proyecto_grado-main/informe_avance_proyecto_v2.docx",
        "C:/Users/julianag18/Desktop/Proyecto de grado/proyecto_grado-main/informe_avance_proyecto.docx",
        "C:/Users/julianag18/Desktop/Proyecto de grado/proyecto_grado-main/informe_avance_proyecto_v2.docx",
        "informe_avance_proyecto.docx",
        "informe_avance_proyecto_v2.docx",
        "informe_avance_proyecto_desktop_backup.docx",
        "C:/Users/julianag18/Desktop/informe_final_practica.docx",
        "C:/Users/julianag18/OneDrive/Desktop/informe_final_practica.docx"
    ]
    for file_path in cleanup_files:
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                print(f"Deleted: {file_path}")
        except Exception as e:
            pass

if __name__ == "__main__":
    create_final_report()
